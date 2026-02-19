"""DOCX document extractor using python-docx."""

from dataclasses import field
from io import BytesIO
from typing import Any

from app.core.document_processing.base import (
    BaseExtractor,
    DocumentType,
    ExtractionError,
    ExtractionResult,
    ExtractedSection,
    ExtractorRegistry,
    SIZE_LIMITS,
)
from app.core.logging import get_logger

logger = get_logger(__name__)


class DOCXExtractor(BaseExtractor):
    """DOCX document extractor.

    Extracts text from Word documents using python-docx.
    Preserves heading structure as semantic sections.
    """

    def can_handle(self, mime_type: str, file_extension: str) -> bool:
        return (
            mime_type
            in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            )
            or file_extension.lower() in (".docx", ".doc")
        )

    def get_supported_types(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]

    def get_size_limit(self) -> int:
        return SIZE_LIMITS.get(DocumentType.DOCX, 10 * 1024 * 1024)

    async def extract(
        self,
        file_bytes: bytes,
        filename: str,
        max_pages: int | None = None,
        **kwargs: Any,
    ) -> ExtractionResult:
        """Extract content from DOCX.

        Parses paragraphs and uses heading styles to create semantic sections.
        Tables are extracted as pipe-delimited text.
        """
        valid, error_msg = self.validate_size(file_bytes)
        if not valid:
            raise ExtractionError(error_msg, extractor="docx", recoverable=False)

        try:
            from docx import Document
        except ImportError:
            raise ExtractionError(
                "python-docx not installed", extractor="docx", recoverable=False
            )

        try:
            doc = Document(BytesIO(file_bytes))
        except Exception as e:
            raise ExtractionError(
                f"Failed to open DOCX: {e}", extractor="docx", recoverable=False
            )

        sections: list[ExtractedSection] = []
        total_words = 0
        all_text_parts: list[str] = []
        warnings: list[str] = []

        # Build sections from headings + body paragraphs
        current_heading: str | None = None
        current_content: list[str] = []
        section_idx = 0

        def _flush_section() -> None:
            nonlocal section_idx, total_words
            text = "\n".join(current_content).strip()
            if not text:
                return
            wc = len(text.split())
            total_words += wc
            all_text_parts.append(text)
            sections.append(
                ExtractedSection(
                    content=text,
                    section_type="heading" if current_heading else "body",
                    section_title=current_heading or f"Section {section_idx + 1}",
                    page_number=section_idx + 1,
                    word_count=wc,
                )
            )
            section_idx += 1

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower() if para.style else ""
            text = para.text.strip()

            if "heading" in style_name and text:
                # Flush previous section
                _flush_section()
                current_heading = text
                current_content = []
            elif text:
                current_content.append(text)

        # Flush final section
        _flush_section()

        # Extract tables
        for idx, table in enumerate(doc.tables):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text.strip():
                wc = len(table_text.split())
                total_words += wc
                all_text_parts.append(table_text)
                sections.append(
                    ExtractedSection(
                        content=table_text,
                        section_type="table",
                        section_title=f"Table {idx + 1}",
                        page_number=0,
                        word_count=wc,
                    )
                )

        raw_text = "\n\n".join(all_text_parts)

        if not raw_text.strip():
            warnings.append("Document appears to be empty or image-only")

        logger.info(
            f"Extracted {len(sections)} sections, {total_words} words from {filename}"
        )

        return ExtractionResult(
            sections=sections,
            page_count=len(sections),
            word_count=total_words,
            extraction_method="native",
            raw_text=raw_text[:200000],
            warnings=warnings,
        )


# Register extractor
ExtractorRegistry.register(DOCXExtractor())
