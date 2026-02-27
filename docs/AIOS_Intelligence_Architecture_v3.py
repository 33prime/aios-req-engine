"""Generate the AIOS Intelligence Architecture v3 DOCX document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x0A, 0x1E, 0x2F)  # Navy
    heading_style.font.name = 'Calibri'

GREEN = RGBColor(0x3F, 0xAF, 0x7A)
NAVY = RGBColor(0x0A, 0x1E, 0x2F)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x66, 0x66, 0x66)


def add_callout(text, label=""):
    """Add a styled callout box (using table with shading)."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(0.4)
    table.columns[1].width = Inches(5.6)

    left_cell = table.cell(0, 0)
    left_cell.text = ""
    # Shade left cell green
    from docx.oxml.ns import qn
    tc = left_cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): '3FAF7A',
        qn('w:val'): 'clear',
    })
    tcPr.append(shading)

    right_cell = table.cell(0, 1)
    if label:
        p = right_cell.paragraphs[0]
        run = p.add_run(label + "\n")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
        run2.font.color.rgb = DARK
    else:
        right_cell.text = text
        for p in right_cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

    doc.add_paragraph("")


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")
    return table


# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_paragraph("")
doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("AIOS")
run.font.size = Pt(42)
run.font.color.rgb = NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("AI Operating System")
run.font.size = Pt(20)
run.font.color.rgb = GREEN

doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run("Intelligence Architecture & Technical Deep Dive")
run.font.size = Pt(16)
run.font.color.rgb = NAVY

p = doc.add_paragraph()
run = p.add_run("ReadyToGo.ai  \u2022  February 2026  \u2022  Confidential")
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_paragraph("")

add_callout(
    "5x Faster (32s vs 206s extraction)  \u2022  100% Coverage (zero document truncation)  \u2022  "
    "7x Cheaper ($0.003 vs $0.021 per doc)  \u2022  $0.04 Total Cost (full pipeline per document)",
    "Performance at a Glance"
)

doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run(
    "The architecture that turns unstructured discovery conversations into defensible, "
    "evidence-grounded business requirements \u2014 automatically."
)
run.font.size = Pt(12)
run.font.color.rgb = NAVY
run.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. THE PROBLEM
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("1. The Problem", level=1)

doc.add_paragraph(
    "Consultants sell expertise, but they deliver documents. The gap between what a client "
    "says in a discovery session and what appears in a Business Requirements Document is "
    "where projects fail. Critical details get lost. Stakeholder context gets flattened. "
    "Decisions made in conversation never make it to the spec."
)

doc.add_paragraph(
    "A single 60-minute discovery session produces 8,000\u201312,000 words of transcript. A "
    "consultant reads it once, extracts what they remember, and discards the rest. The typical "
    "workflow \u2014 record, manually review, copy-paste into a template, cross-reference with "
    "previous sessions \u2014 is slow, lossy, and unscalable."
)

doc.add_paragraph(
    "AIOS eliminates this gap entirely. Every paragraph of every conversation is semantically "
    "indexed, structurally tagged, and traceable through the final deliverable. When the BRD "
    'says "the client prioritizes voice-first UX," you can click through to the exact transcript '
    "paragraph where the CEO said it, at timestamp 14:32, in the context of discussing mobile workflows."
)

add_callout(
    "The value isn\u2019t in generating documents \u2014 any LLM can do that.\n"
    "The value is in provenance: every claim traces to evidence, every entity traces to a source, "
    "every recommendation traces to a stakeholder\u2019s actual words. And in retrieval: the system "
    "doesn\u2019t just store knowledge \u2014 it finds the right knowledge for every question, automatically.",
    "The Core Insight"
)

# ═══════════════════════════════════════════════════════════════════════════
# 2. THREE-LAYER ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("2. The Three-Layer Intelligence Architecture", level=1)

doc.add_paragraph(
    "AIOS is organized into three distinct layers, connected by a unified retrieval system "
    "that makes every flow in the platform evidence-aware."
)

add_callout(
    "Most AI systems connect Layer 1 directly to Layer 3 \u2014 raw input to structured output. "
    "They truncate, they summarize, they lose context. AIOS has a retrieval system that sits "
    "between the layers, making every flow in the platform evidence-aware. Any component that "
    "needs context \u2014 chat, briefing, solution design, gap analysis \u2014 asks one question and gets "
    "chunks + entities + beliefs + graph relationships, ranked by relevance.",
    "Connecting the Layers: Unified Retrieval"
)

doc.add_heading("Layer 1: Signals (Raw Input)", level=2)
doc.add_paragraph(
    "Signals are everything the system ingests: meeting transcripts, uploaded documents, emails, "
    "chat messages, client portal responses, and prototype session feedback. Each signal is a "
    "raw, unprocessed artifact from the real world, arriving through multiple channels \u2014 direct "
    "upload, email forwarding, portal submission, or live session capture."
)

doc.add_heading("Layer 2: Intelligence (Understanding)", level=2)
doc.add_paragraph(
    "This is where AIOS transforms raw input into structured, searchable knowledge. "
    "Five subsystems work together:"
)

doc.add_paragraph(
    "Semantic Chunks \u2014 Documents are split into sections at natural boundaries (headings, "
    "speaker turns, topic shifts). Each chunk is a self-contained unit of meaning with metadata: "
    "section title, speaker, page number, temporal context."
)

doc.add_paragraph(
    "Vector Embeddings \u2014 Every chunk is converted to a 1,536-dimension vector using "
    "OpenAI\u2019s text-embedding-3-small. This vector is a semantic fingerprint that captures "
    'what the chunk is "about" regardless of specific word choice. Searching by vector '
    "similarity finds relevant content even when different vocabulary is used."
)

doc.add_paragraph(
    "Meta-Tags \u2014 A lightweight Haiku pass ($0.001 per chunk) enriches each chunk with "
    "structured metadata: entities mentioned, topics discussed, decisions made, speaker roles, "
    "temporal context, and confidence signals. These tags enable hybrid search \u2014 combining "
    "semantic similarity with structured filtering \u2014 without a separate full-text search index."
)

doc.add_paragraph(
    "Entity Embeddings \u2014 Every BRD entity (features, workflows, constraints, stakeholders, "
    "etc.) is embedded as a vector alongside signal chunks. This creates a unified search layer: "
    "one query returns both the transcript paragraphs where the client discussed payments AND "
    "the features, workflows, and data entities that address payments."
)

doc.add_paragraph(
    "Memory Graph \u2014 A belief network that synthesizes facts extracted from chunks into "
    "higher-order beliefs, contradictions, tensions, and strategic insights. Each belief traces "
    "to the facts that support it, which trace to the chunks they were extracted from, which "
    "trace to the original signal. Four levels of abstraction, full provenance at every level, "
    "with embedded vectors for semantic contradiction detection."
)

doc.add_heading("Layer 3: BRD (Structured Output)", level=2)
doc.add_paragraph(
    "The business requirements document is a living knowledge graph \u2014 features, workflows, "
    "business drivers, constraints, stakeholders, data entities, competitors, assumptions, and "
    "their relationships. Every entity has a provenance chain back through the intelligence "
    "layer to the original signal. The BRD evolves as new signals arrive, beliefs are "
    "confirmed or contradicted, and requirements are refined through iterative discovery."
)

# ═══════════════════════════════════════════════════════════════════════════
# 3. SIGNAL PROCESSING PIPELINE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("3. The Signal Processing Pipeline", level=1)

doc.add_paragraph(
    "When a document enters AIOS, it goes through a multi-stage pipeline implemented as a "
    "LangGraph state machine with conditional routing. Each stage\u2019s output determines the "
    "next stage\u2019s behavior."
)

doc.add_heading("Stage 1: Ingestion & Triage", level=2)
doc.add_paragraph(
    "The system classifies the signal by type (transcript, document, email, feedback) and "
    "assesses complexity. Short signals take a fast path. Long documents proceed to full "
    "processing. Pure heuristic \u2014 no LLM call, under 100ms. The classifier also generates "
    "keyword tags and topic labels that propagate to every chunk downstream at zero additional cost."
)

doc.add_heading("Stage 2: Semantic Chunking", level=2)
doc.add_paragraph(
    "Documents are split at natural boundaries: heading changes, speaker turns in transcripts, "
    "topic shifts detected by embedding discontinuity. Chunks carry metadata: section title, "
    "page or timestamp reference, position index. This is structure-aware splitting that "
    "respects the document\u2019s own organization."
)

doc.add_heading("Stage 3: Meta-Tag Enrichment (Parallel Haiku)", level=2)
doc.add_paragraph(
    "Each chunk receives a parallel Haiku call that produces structured metadata tags. These "
    "tags serve as a pre-computed index enabling hybrid search at query time without additional "
    "LLM calls:"
)

add_table(
    ["Tag", "Purpose", "Query-Time Benefit"],
    [
        ["entities_mentioned", "Who and what appears in this chunk", 'Filter: "chunks mentioning Sarah Chen"'],
        ["topics", "Semantic topic labels (snake_case slugs)", 'Filter + boost: "chunks about payment_processing"'],
        ["decision_made", "Boolean: was something decided here", "Filter: only confirmed decisions"],
        ["speaker_roles", "Maps speakers to org roles", "Speaker \u2192 stakeholder resolution without LLM"],
        ["temporal", "current_state vs. future_state", "Filter: pain points vs. aspirations"],
        ["confidence_signals", "explicit_requirement, preference, speculation", "Weight: hard requirements rank higher"],
        ["entity_types_discussed", "Which BRD entity types appear", "Filter by feature, workflow, constraint, etc."],
    ]
)

doc.add_heading("Stage 4: Vector Embedding", level=2)
doc.add_paragraph(
    "Each chunk is embedded using OpenAI\u2019s text-embedding-3-small (1,536 dimensions) with "
    'a context-aware prefix prepended \u2014 "Discovery transcript, Section: Voice Input '
    'Preferences" \u2014 so the vector captures both content and structural context. Meta-tagging '
    "and embedding run in parallel (both are independent of each other), adding zero sequential latency."
)

doc.add_heading("Stage 5: Parallel Entity Extraction (Haiku Map-Reduce)", level=2)
doc.add_paragraph(
    "Rather than sending a truncated document to a single expensive model, AIOS fans out one "
    "lightweight Haiku call per chunk, processes all chunks simultaneously, then merges and "
    "deduplicates the results. Each Haiku call receives the full entity inventory (so it knows "
    "what already exists), but only sees one section."
)

doc.add_paragraph(
    "The static system prompt uses Anthropic\u2019s prompt caching \u2014 all parallel calls share "
    "the same cached prefix, paying only for cache reads after the first call. For a document "
    "with 11 chunks, 10 of the 11 calls hit the cache at 90% discount."
)

add_callout(
    "Before: Single Sonnet call, 206 seconds, 12K char truncation (59% discarded), 44 patches.\n"
    "After: 11 parallel Haiku calls, 32 seconds, zero truncation, 108 patches extracted.\n"
    "Result: 5x faster, 7x cheaper, 2.5x more entities discovered, zero information loss.",
    "Measured Performance"
)

doc.add_heading("Stage 6: Four-Tier Entity Deduplication", level=2)
doc.add_paragraph(
    "When multiple chunks mention the same entity, parallel extraction produces duplicate "
    "patches. The deduplication pipeline resolves these in four tiers with per-entity-type "
    "thresholds, each catching what the previous tier missed:"
)

add_table(
    ["Tier", "Method", "Speed", "Cost", "Catches"],
    [
        ["1. Exact", "Normalized string match", "<1ms", "Free",
         "Identical names with different casing or whitespace"],
        ["2. Fuzzy", "RapidFuzz (token-set + partial + weighted)", "<1ms", "Free",
         'Rewordings: "Push Notifications" \u2194 "Notification System"'],
        ["2.5 Semantic Rerank", "Cohere rerank-v3.5 relevance scoring", "~20ms", "$0.00002",
         '"Reduce assessment time by 40%" \u2194 "Cut evaluation duration in half"'],
        ["3. Embedding", "Cosine similarity via pgvector", "~50ms", "$0.0001",
         'Meaning equivalence: "HIPAA Required" \u2194 "HIPAA Compliance"'],
    ]
)

doc.add_paragraph(
    "Tier 2.5 fires only in the ambiguous zone \u2014 when the fuzzy score falls between the "
    "ambiguous threshold and the merge threshold. Cohere\u2019s rerank-v3.5 is purpose-built for "
    "semantic relevance scoring: it reads the new entity name alongside every existing entity "
    "of the same type and returns a relevance score. If the score exceeds 0.8, it\u2019s a match. "
    "If Cohere is unavailable, the system silently falls through to the existing embedding tier."
)

doc.add_paragraph(
    "The result: 80\u201390% of dedup decisions still resolve instantly at zero cost (tiers 1\u20132). "
    "Tier 2.5 catches the subtle semantic duplicates that string matching misses \u2014 "
    "different phrasing of the same requirement \u2014 before falling back to the heavier "
    "embedding comparison. Per-type configuration prevents false merges: competitors with "
    "short names skip embedding entirely; business drivers with generic descriptions like "
    '"increase revenue" are rejected before any matching runs.'
)

doc.add_heading("Stage 7: Speaker Resolution", level=2)
doc.add_paragraph(
    "Meta-tags from Stage 3 include speaker_roles \u2014 a mapping of speaker names to "
    'organizational roles ("Bran" \u2192 "client_ceo"). A three-tier fuzzy resolver links these '
    "to stakeholder entities: exact first/last name match, fuzzy name match (token-set ratio "
    '> 0.75), and initial+last name match ("B. Wilson" \u2192 "Brandon Wilson"). This runs as '
    "post-processing after patch application and automatically increments topic_mentions "
    "counters on stakeholder records."
)

doc.add_heading("Stage 8: Memory Synthesis", level=2)
doc.add_paragraph(
    "The Memory Watcher processes each chunk in parallel (same fan-out pattern as extraction). "
    "Each Watcher call reads one chunk plus the 5 most relevant existing beliefs (retrieved by "
    "embedding similarity), extracting grounded facts with direct quotes, chunk IDs, and "
    "speaker attribution. Facts with importance \u2265 0.7 or detected contradictions trigger the "
    "Synthesizer, which updates the belief network \u2014 creating beliefs, adjusting confidence, "
    "adding supporting or contradicting edges. Contradiction detection uses vector similarity "
    "between new facts and existing beliefs, catching semantic conflicts even when vocabulary "
    "differs completely."
)

doc.add_heading("Stage 9: Question Auto-Resolution", level=2)
doc.add_paragraph(
    "After entities are saved and memory is updated, the system re-evaluates open questions. "
    "Every open question \u2014 both project-level and per-solution-flow-step \u2014 is vector-compared "
    "against the new signal\u2019s chunks. A lightweight Haiku call checks whether the new evidence "
    "answers the question with confidence above 0.80. If so, the question is auto-resolved "
    "with a citation linking back to the chunk. This same mechanism fires when clients answer "
    "portal questions and when confirmation clusters are approved \u2014 every new piece of "
    "evidence is checked against every open question."
)

# ═══════════════════════════════════════════════════════════════════════════
# 4. PULSE ENGINE (NEW SECTION)
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("4. The Pulse Engine: Adaptive Project Intelligence", level=1)

doc.add_paragraph(
    "The Pulse Engine is a zero-LLM deterministic rules engine that computes project health "
    "in ~50ms and adapts system behavior based on what\u2019s actually happening. It fires "
    "after every state-changing event \u2014 signal processed, entity confirmed, cluster merged \u2014 "
    "and produces a comprehensive health snapshot that other systems consume."
)

doc.add_heading("What It Computes", level=2)

doc.add_paragraph(
    "For every entity type (features, personas, workflows, constraints, etc.), the engine "
    "computes real-time health metrics:"
)

add_table(
    ["Metric", "What It Measures", "Example"],
    [
        ["Coverage", "Count vs. stage-appropriate target",
         "8 features vs. target of 8 = adequate; 2 personas vs. target of 3 = thin"],
        ["Confirmation Rate", "% of entities confirmed by consultant or client",
         "6/8 features confirmed = 75%"],
        ["Staleness Rate", "% of entities marked stale (contradicted by new evidence)",
         "1/8 features stale = 12.5%"],
        ["Quality Score", "Weighted blend: 60% confirmation + 40% freshness",
         "0.75 \u00d7 0.6 + 0.875 \u00d7 0.4 = 0.80"],
        ["Health Score", "Stage-weighted composite (0\u2013100)",
         "Discovery weights coverage 50%; Validation weights confirmation 45%"],
    ]
)

doc.add_heading("Stage-Aware Intelligence", level=2)
doc.add_paragraph(
    "The engine knows which project stage matters and adjusts priorities accordingly. "
    "In Discovery, coverage dominates (50% weight) \u2014 the system needs entities to exist. "
    "In Validation, confirmation dominates (45% weight) \u2014 the system needs entities to be "
    "verified. Transition gates are hard rules: you can\u2019t leave Discovery until you have "
    "\u22655 features, \u22652 personas, \u22652 workflows. You can\u2019t leave Validation until "
    "\u22654 confirmed features and \u22652 confirmed personas."
)

doc.add_heading("Adaptive Behavior: How the Pulse Changes the System", level=2)

p = doc.add_paragraph()
run = p.add_run("Signal Velocity Scaling. ")
run.bold = True
p.add_run(
    "The engine compares signal volume in the first half vs. second half of a 7-day window. "
    "If signals are accelerating (2nd half \u2265 1.5\u00d7 1st half), entity targets scale up 20% \u2014 "
    "more signals means more entities expected. If signals are stalling (2nd half \u2264 0.5\u00d7), "
    "targets scale down 15% to avoid false \"thin coverage\" warnings."
)

p = doc.add_paragraph()
run = p.add_run("Dynamic Dedup Thresholds. ")
run.bold = True
p.add_run(
    "This is where it gets interesting. The pulse health map feeds directly into entity "
    "deduplication. When a type is saturated (\u2265100% of target), the fuzzy merge threshold "
    "drops by 0.05 \u2014 merge more aggressively, because adding duplicates hurts more than "
    "missing a new entity. When a type is missing or thin (<30% of target), the threshold "
    "rises by 0.05 \u2014 be lenient, because creating new entities helps more than a false merge."
)

p = doc.add_paragraph()
run = p.add_run("Extraction Directives. ")
run.bold = True
p.add_run(
    "Each entity type receives a directive that shapes what the extraction pipeline does: "
    "GROW (create freely), CONFIRM (stop creating, start confirming), ENRICH (add detail to "
    "existing), MERGE_ONLY (no new entities, only merge duplicates), or STABLE (minimal changes). "
    "These directives propagate into the extraction prompt itself."
)

p = doc.add_paragraph()
run = p.add_run("Action Ranking. ")
run.bold = True
p.add_run(
    "The engine ranks next actions by impact score (0\u2013100). Actions that would unblock a "
    "stage transition gate get a 2\u00d7 multiplier. The top 5 actions surface in the chat "
    "assistant\u2019s context frame and the home dashboard\u2019s \"Next Steps\" cards."
)

add_callout(
    "The Pulse Engine means AIOS gets smarter as projects progress. A new project with 3 "
    "features gets aggressive entity creation. A mature project with 12 confirmed features "
    "gets aggressive deduplication. The thresholds, targets, and priorities all adapt to "
    "where the project actually is \u2014 not where a static config assumes it should be.",
    "Why This Matters"
)

# ═══════════════════════════════════════════════════════════════════════════
# 5. UNIFIED RETRIEVAL SYSTEM
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("5. The Unified Retrieval System", level=1)

doc.add_paragraph(
    "Every flow in AIOS \u2014 chat, briefing, solution design, gap analysis, stakeholder "
    "profiling, unlock generation \u2014 calls a single retrieval function. This function "
    "combines six retrieval stages in a pipeline that adapts to query complexity:"
)

doc.add_heading("Stage 1: Agentic Decomposition", level=2)
doc.add_paragraph(
    'A lightweight Haiku call decomposes the query into 2\u20134 targeted sub-queries. "What '
    'are the main risks if we prioritize voice-first UX?" becomes three focused searches: '
    '"voice input technical constraints," "voice UX accessibility concerns," and "timeline '
    'budget risk pressure." For simple, specific queries, the decomposition step is skipped '
    "entirely \u2014 the system detects this automatically."
)
doc.add_paragraph(
    "This solves the fundamental limitation of single-query vector search: no single chunk "
    "answers a multi-faceted question. The decomposition ensures each facet gets its own "
    "targeted search."
)

doc.add_heading("Stage 2: Parallel Retrieval (Three Strategies)", level=2)
doc.add_paragraph(
    "Three retrieval strategies execute simultaneously via asyncio.gather(), each finding "
    "what the others miss:"
)

p = doc.add_paragraph()
run = p.add_run("2a. Vector Search \u2014 ")
run.bold = True
p.add_run(
    "For each sub-query, cosine similarity search on signal_chunks with optional JSONB "
    "meta-tag filters. Finds semantically relevant content: \"payment processing\" matches "
    '"invoice reconciliation" even with zero word overlap. Results deduplicated across sub-queries.'
)

p = doc.add_paragraph()
run = p.add_run("2b. Entity Embedding Search \u2014 ")
run.bold = True
p.add_run(
    "Vector search across all entity embeddings (features, workflows, personas, constraints, "
    "etc.) filtered by page context. When you\u2019re on the Workflows page, entity search returns "
    "only workflows and workflow steps. When you\u2019re on the Solution Flow, it returns flow "
    "steps, features, workflows, and unlocks. Reverse provenance is the fallback: if vector "
    "search finds no entities, the system looks up which entities were created from the chunks "
    "found in 2a."
)

p = doc.add_paragraph()
run = p.add_run("2c. Memory Belief Lookup \u2014 ")
run.bold = True
p.add_run(
    "Vector search on embedded memory_nodes for beliefs relevant to the query. Returns beliefs "
    "with confidence scores and their supporting facts. Provides the system\u2019s current "
    "understanding alongside the raw evidence."
)

doc.add_heading("Stage 2.5: Graph Expansion (Entity Neighborhood Traversal)", level=2)
doc.add_paragraph(
    "This is where retrieval goes beyond vector similarity. The top 3 entities from Stage 2 "
    "become seeds for graph expansion. For each seed, the system queries its 1-hop "
    "neighborhood via signal_impact co-occurrence: which other entities share evidence chunks "
    "with this entity?"
)

doc.add_paragraph(
    "Concrete example: A query about \"payment processing\" returns the Payments feature from "
    "Stage 2. Graph expansion then pulls in the Invoice Data Entity (shares 4 chunks), the "
    "HIPAA Compliance constraint (shares 3 chunks), the Finance Manager persona (shares 2 "
    "chunks), and the Checkout Workflow (shares 2 chunks). The query never mentioned any of "
    "these \u2014 the graph found them through structural relationships."
)

doc.add_paragraph(
    "Three parallel Supabase lookups run in ~50ms total. Up to 15 graph-expanded entities "
    "are added, each marked with source=\"graph_expansion\" so the reranker can properly "
    "rank them against vector-found results. Evidence chunks from neighborhoods are also "
    "merged into the result set."
)

add_callout(
    "Graph expansion solves a critical blind spot in pure vector search: structural "
    "relationships that aren\u2019t captured by semantic similarity. A constraint that affects "
    "a feature doesn\u2019t need to use similar words \u2014 they\u2019re connected because they were "
    "discussed in the same evidence. The graph captures this.",
    "Why Graph + Vector"
)

doc.add_heading("Stage 3: Cohere Reranking (with Haiku Fallback)", level=2)
doc.add_paragraph(
    "After parallel retrieval and graph expansion return 15\u201330 candidates, a reranking pass "
    "scores each candidate against the original query for precision. The system uses a "
    "three-tier reranking strategy:"
)

p = doc.add_paragraph()
run = p.add_run("Primary: Cohere rerank-v3.5 \u2014 ")
run.bold = True
p.add_run(
    "A purpose-built reranking model that reads the query and each chunk together and "
    "returns a relevance score. Up to 25 chunks scored in a single API call (~100ms). "
    "This is faster, cheaper, and more accurate than using a generalist LLM for reranking."
)

p = doc.add_paragraph()
run = p.add_run("Fallback: Haiku Listwise \u2014 ")
run.bold = True
p.add_run(
    "If Cohere is unavailable (no API key or API error), a Haiku call ranks chunks by "
    "relevance in a single prompt, returning an ordered list of the top-k chunk numbers."
)

p = doc.add_paragraph()
run = p.add_run("Final Fallback: Cosine Order \u2014 ")
run.bold = True
p.add_run(
    "If both rerankers fail, chunks are simply truncated by their original cosine similarity "
    "order. The system never fails \u2014 it degrades gracefully."
)

doc.add_heading("Stage 4: Evaluate & Loop", level=2)
doc.add_paragraph(
    "A Haiku evaluation checks whether the retrieved results actually answer the query. If "
    "not, it identifies what\u2019s missing and generates reformulated sub-queries for another "
    "retrieval round. The evaluation criteria are flow-specific: gap intelligence needs "
    '"evidence confirming or denying the gap," solution design needs "pain + goals + '
    'constraints," briefing needs "supporting AND contradicting evidence."'
)
doc.add_paragraph(
    "This agentic loop fires at most 2\u20133 rounds. For 70% of queries, one round is "
    "sufficient. The loop catches cases where the first search returned topically related "
    "but incomplete results."
)

doc.add_heading("Stage 5: Structured Formatting", level=2)
doc.add_paragraph(
    "Results are formatted for injection into any LLM context. Formatting adapts by use case:"
)

add_table(
    ["Style", "Used By", "What It Produces"],
    [
        ["chat", "Chat assistant, stakeholder intel",
         "Concise quotes with entity links, max 1\u20132K tokens"],
        ["generation", "Unlocks, prototype planner, solution flow",
         "Grouped by entity type with full descriptions, 600\u20133K tokens"],
        ["analysis", "Briefing engine, gap intelligence",
         "Supporting vs. contradicting evidence with confidence scores, max 1.5K tokens"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════
# 6. TAILORED CONTEXT PER PIPELINE (NEW SECTION)
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("6. Tailored Context: Every Pipeline Gets What It Needs", level=1)

doc.add_paragraph(
    "Every pipeline in AIOS calls the same retrieve() function \u2014 but with dramatically "
    "different parameters. The system doesn\u2019t just retrieve generically. It knows what each "
    "pipeline needs and tunes retrieval accordingly."
)

doc.add_heading("How Different Pipelines Call Retrieval", level=2)

add_table(
    ["Pipeline", "Query Shape", "Rounds", "Entity Filter", "Evidence Style", "Key Insight"],
    [
        ["Chat (simple question)", "User\u2019s message verbatim", "1",
         "Filtered by current page", "chat (concise)",
         "Simple questions skip decomposition AND reranking for instant response"],
        ["Chat (complex question)", "User\u2019s message", "2",
         "Filtered by current page", "chat",
         "Haiku decomposes then retrieves; Cohere reranks"],
        ["Briefing Engine", '"evidence for and against: {tension}"', "2",
         "All types", "analysis (pro/con)",
         "Explicitly requests BOTH supporting and contradicting evidence"],
        ["Gap Intelligence", '"gaps uncertainties unknowns in {phase}"', "1",
         "All types", "analysis",
         "All skip flags ON \u2014 pure vector search, no LLM overhead"],
        ["Unlocks (3 parallel)", "Tier-specific: operational / strategic / visionary", "1",
         "Tier-specific filters", "generation",
         "3 concurrent calls with different entity filters per unlock tier"],
        ["Prototype Planner", '"feature names + feedback summary"', "2",
         "All types", "generation (3K tokens)",
         'Criteria: "Need original requirements evidence for these features"'],
        ["Stakeholder Intel", '"person name + role + concerns"', "1",
         "All types", "chat",
         "All skip flags ON \u2014 fast lookup of who said what about whom"],
    ]
)

doc.add_heading("Page-Aware Entity Filtering", level=2)
doc.add_paragraph(
    "When the consultant is viewing the Workflows page, the chat assistant only retrieves "
    "workflow and workflow_step entities. On the Features page, features and unlocks. On "
    "the Solution Flow, solution_flow_steps, features, workflows, and unlocks. This filtering "
    "happens via a 10-page mapping that narrows entity search from all 12 types to the "
    "3\u20134 most relevant per page."
)

add_table(
    ["Page Context", "Entity Types Retrieved"],
    [
        ["brd:features", "feature, unlock"],
        ["brd:personas", "persona"],
        ["brd:workflows", "workflow, workflow_step"],
        ["brd:constraints", "constraint"],
        ["brd:solution-flow", "solution_flow_step, feature, workflow, unlock"],
        ["brd:unlocks", "unlock, feature, competitor"],
        ["brd:business-drivers", "business_driver"],
        ["brd:stakeholders", "stakeholder"],
        ["prototype", "prototype_feedback, feature"],
        ["overview / canvas", "All types (no filter)"],
    ]
)

doc.add_heading("Solution Flow: Zero-LLM Context Layers", level=2)
doc.add_paragraph(
    "When a consultant is working on the Solution Flow \u2014 the step-by-step path from pain "
    "point to solution \u2014 the system builds 4 layers of context without a single LLM call:"
)

p = doc.add_paragraph()
run = p.add_run("Layer 1: Flow Summary. ")
run.bold = True
p.add_run(
    "~50 tokens per step. Phase, title, actors, field counts. The model sees the full flow "
    "structure at a glance."
)

p = doc.add_paragraph()
run = p.add_run("Layer 2: Focused Step Detail. ")
run.bold = True
p.add_run(
    "~300 tokens. The currently-viewed step\u2019s goal, actors, fields, open questions, with "
    "batch-resolved entity names (features, workflows, data entities referenced by this step)."
)

p = doc.add_paragraph()
run = p.add_run("Layer 3: Cross-Step Intelligence. ")
run.bold = True
p.add_run(
    "~200 tokens. Six deterministic checks across all steps: actor gaps, phase coverage, "
    "confidence distribution, explore hotspots, data flow analysis, staleness detection."
)

p = doc.add_paragraph()
run = p.add_run("Layer 4: Retrieval Hints. ")
run.bold = True
p.add_run(
    "4 strings derived from the focused step\u2019s goal, open questions, and actors. These "
    "become the context_hint for query decomposition, steering vector search toward relevant evidence."
)

doc.add_paragraph(
    "All four layers are built from database queries and string formatting. The result is "
    "~750 tokens of rich, structured context injected into the prompt at near-zero latency "
    "and zero cost."
)

# ═══════════════════════════════════════════════════════════════════════════
# 7. PROMPT CACHING (NEW SECTION)
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("7. Prompt Caching: Pay Once, Read Forever", level=1)

doc.add_paragraph(
    "AIOS uses Anthropic\u2019s prompt caching across 20+ LLM call sites. The pattern is "
    "consistent: expensive system prompts and stable project context are cached for 5 minutes. "
    "Subsequent calls within that window pay only 10% of the input token cost."
)

doc.add_heading("The Parallel Fan-Out Pattern", level=2)
doc.add_paragraph(
    "This is where prompt caching saves the most. Consider entity extraction: 11 chunks each "
    "get their own Haiku call, but all 11 share the same system prompt (extraction rules, "
    "entity type definitions, entity inventory). The first call creates the cache. The remaining "
    "10 calls read from cache at 90% discount."
)

doc.add_paragraph("The same pattern applies to:")

p = doc.add_paragraph()
run = p.add_run("Meta-tagging: ")
run.bold = True
p.add_run(
    "N chunks \u00d7 1 Haiku call each. Same cached system prompt (tag definitions, output schema). "
    "All parallel."
)

p = doc.add_paragraph()
run = p.add_run("Memory observation: ")
run.bold = True
p.add_run(
    "N chunks \u00d7 1 Haiku call each. Same cached system prompt (fact extraction rules, belief format)."
)

p = doc.add_paragraph()
run = p.add_run("Chat assistant: ")
run.bold = True
p.add_run(
    "Two-block architecture. Block 1 (static identity, capabilities, action card schemas, "
    "tool decision tree) is cached. Block 2 (per-request project state, gaps, focused entity) "
    "is dynamic. Tool definitions are also cached. Across a 5-turn conversation, the static "
    "block is created once and read 4 times."
)

doc.add_heading("What Gets Cached", level=2)

add_table(
    ["Component", "Cached Content", "Approx. Tokens", "Pattern"],
    [
        ["Entity extraction", "Extraction rules + entity inventory", "~2,500", "11 parallel calls"],
        ["Meta-tagging", "Tag definitions + output schema", "~400", "N parallel calls"],
        ["Chat assistant", "Identity + capabilities + tool tree + action cards", "~3,000", "Multi-turn"],
        ["Solution flow gen", "Generation rules + project context", "~2,000 + context", "Iterative retries"],
        ["Briefing narrative", "Tone, style, output format", "~1,200", "Per-briefing"],
        ["Gap intelligence", "Gap identification rules", "~1,500", "Per-analysis"],
        ["Intelligence agents", "Agent system prompt (SI, CI, Research)", "~3,000", "Multi-turn tool loop"],
    ]
)

add_callout(
    "For a typical document with 11 chunks: meta-tagging saves ~4K cached tokens across 10 "
    "reads. Extraction saves ~25K cached tokens across 10 reads. Memory saves ~4K across 10 "
    "reads. Total: ~33K tokens at 90% discount per document processed. Across a project with "
    "50 signals, that\u2019s ~1.65M tokens saved \u2014 roughly $2.50 in avoided Haiku input costs alone.",
    "Cost Impact"
)

# ═══════════════════════════════════════════════════════════════════════════
# 8. KNOWLEDGE GRAPH INTELLIGENCE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("8. Knowledge Graph Intelligence", level=1)

doc.add_paragraph(
    "The BRD entity tables with their foreign keys and the signal_impact provenance table "
    "form an implicit knowledge graph. AIOS queries this graph for structural intelligence "
    "that vector search alone cannot provide."
)

doc.add_heading("Entity Neighborhoods", level=2)
doc.add_paragraph(
    "For any entity, the system retrieves its structural context via the signal_impact table: "
    "what chunks created it, what other entities share those chunks (co-occurrence), what "
    "constraints apply, which stakeholders are connected, what workflows it belongs to. "
    "Co-occurring entities are ranked by shared chunk count \u2014 the more evidence they share, "
    "the stronger the relationship. This is JOIN-based graph traversal on existing tables. "
    "No graph database required."
)

doc.add_heading("Relationship Path Finding", level=2)
doc.add_paragraph(
    'When the chat receives a relationship question \u2014 "How does the payment workflow '
    'connect to Sarah\u2019s HIPAA concerns?" \u2014 the system traces paths through the entity '
    "graph using BFS on the signal_impact table. Payment Workflow \u2192 processes \u2192 Invoice "
    "Data Entity \u2192 constrained by \u2192 HIPAA Compliance \u2192 raised by \u2192 Sarah Chen. "
    "Maximum 3 hops, under 50ms, pure SQL."
)

doc.add_heading("Structural Tension Detection", level=2)
doc.add_paragraph(
    "Domain-specific SQL pattern queries identify risks that no single chunk or entity reveals:"
)

p = doc.add_paragraph()
run = p.add_run("Ungrounded Features \u2014 ")
run.bold = True
p.add_run("Confirmed features with zero evidence chunks in signal_impact.")

p = doc.add_paragraph()
run = p.add_run("Stale Assumptions \u2014 ")
run.bold = True
p.add_run("AI-generated entities unvalidated for over 7 days.")

p = doc.add_paragraph()
run = p.add_run("Conflicting Beliefs \u2014 ")
run.bold = True
p.add_run("Two beliefs with confidence > 0.5 connected by a contradiction edge in the memory graph.")

p = doc.add_paragraph()
run = p.add_run("High Pain, No Solution \u2014 ")
run.bold = True
p.add_run(
    "Manual workflow steps with pain severity \u22654 and no features addressing them."
)

# ═══════════════════════════════════════════════════════════════════════════
# 9. CHAT INTELLIGENCE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("9. Chat Intelligence", level=1)

doc.add_paragraph(
    "The chat assistant is the primary interface for consultants and clients. Its architecture "
    "follows one principle: reads should be instant, writes should be deliberate."
)

doc.add_heading("Vector Pre-Fetch (Zero-Tool Reads)", level=2)
doc.add_paragraph(
    "Before the LLM processes any message, the unified retrieval system embeds the message "
    "and retrieves the most relevant chunks, entities, and beliefs (~200ms). These are "
    "injected into the system prompt as evidence context. The LLM responds in one turn with "
    "grounded answers \u2014 no tool call, no second inference pass."
)
doc.add_paragraph(
    "Approximately 70% of chat interactions are informational queries. Vector pre-fetch "
    "handles all of these in a single LLM turn. Tools are reserved for write operations: "
    "creating entities, updating statuses, generating deliverables."
)

doc.add_heading("Context Frame Caching", level=2)
doc.add_paragraph(
    "A cached context frame per project includes: current phase and completion percentage, "
    "top knowledge gaps, project state summary, workflow structure, and low-confidence beliefs "
    "requiring attention. The frame uses fingerprint-based caching \u2014 it hashes entity counts "
    "and workflow structure into a 16-character digest. If nothing changed, it returns "
    "instantly. Any entity mutation invalidates the cache and triggers a rebuild."
)

doc.add_heading("Page-Aware Tool Filtering", level=2)
doc.add_paragraph(
    "The tool set adapts based on which page the user is viewing. On the Workflows page, "
    "workflow CRUD tools are prioritized with guidance like \"When user says \u2018add step after "
    "X\u2019, find the step in context, compute the next step_number, and create with the "
    'correct workflow_id." On the Features page, feature creation and enrichment are '
    "prioritized. This reduces tool definition token count by 40\u201360% and keeps the model "
    "focused on relevant actions."
)

doc.add_heading("Dynamic Two-Block Prompts", level=2)
doc.add_paragraph(
    "Every chat message gets a two-block system prompt. Block 1 (static) contains identity, "
    "capabilities, action card schemas, tool decision tree, and conversation patterns \u2014 "
    "cached via Anthropic\u2019s prompt cache. Block 2 (dynamic) contains the current project "
    "state, phase progress, active gaps (top 5), workflow context, memory hints, page awareness, "
    "and focused entity details. In a multi-turn conversation, Block 1 is created once and "
    "read from cache for every subsequent turn."
)

# ═══════════════════════════════════════════════════════════════════════════
# 10. MEMORY & BELIEF NETWORK
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("10. Memory & Belief Network", level=1)

doc.add_paragraph(
    "The memory system transforms individual observations into organizational knowledge at "
    "four levels of abstraction:"
)

add_table(
    ["Level", "What It Is", "Source", "Example"],
    [
        ["Evidence", "Raw quotes from chunks with speaker + timestamp",
         "Signal processing", '"I capture my best ideas on walks" \u2014 Bran, 14:32'],
        ["Facts", "Verified observations with chunk_id + source_quote",
         "Watcher (Haiku, per-chunk)", "CEO prefers voice input over text for idea capture"],
        ["Beliefs", "Synthesized conclusions with confidence 0\u20131",
         "Synthesizer (Sonnet)", "Client prioritizes voice-first UX (confidence: 0.87)"],
        ["Insights", "Strategic patterns across beliefs",
         "Reflector (Sonnet)", "Tension: scope ambition vs. team capacity for voice features"],
    ]
)

doc.add_paragraph(
    "Every memory node is embedded as a vector. Contradiction detection computes cosine "
    "distance between new facts and existing beliefs \u2014 catching semantic conflicts even "
    "when vocabulary differs completely. Belief confidence evolves over time: corroborating "
    "evidence from independent signals increases confidence, contradicting evidence decreases "
    "it, client confirmation through the portal pushes confidence above 0.9."
)

# ═══════════════════════════════════════════════════════════════════════════
# 11. CLIENT PORTAL
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("11. Client Portal & Intelligent Information Gathering", level=1)

doc.add_paragraph(
    "The client portal is an intelligent information gathering system that knows what it "
    "doesn\u2019t know and figures out the best way to find out."
)

doc.add_heading("Gap \u2192 Source \u2192 Channel Routing", level=2)
doc.add_paragraph(
    "When the system identifies a knowledge gap, it traces through the intelligence layer: "
    "vector search finds who discussed similar topics, chunk metadata reveals the speaker, "
    "speaker resolution links to the stakeholder. The system then assesses complexity using "
    "belief confidence \u2014 high-confidence beliefs need email confirmation, low-confidence "
    "unknowns need meeting discussion, true gaps need portal discovery questions."
)

doc.add_heading("Confirmation Clustering", level=2)
doc.add_paragraph(
    "Individual confirmation items are embedded and clustered by cosine similarity (threshold "
    "0.78) before routing. Three items about budget, timeline, and team allocation cluster as "
    "one conversation about project scoping. The system presents grouped topics rather than "
    "atomic questions."
)

doc.add_heading("Dynamic Question Management", level=2)
doc.add_paragraph(
    "When a client answers a portal question, the answer enters the full signal pipeline. "
    "The system re-evaluates remaining questions \u2014 vector-comparing the answer against open "
    "question embeddings \u2014 and auto-resolves redundancies above 0.80 confidence. Partial "
    "answers accumulate evidence until a threshold is crossed."
)

# ═══════════════════════════════════════════════════════════════════════════
# 12. PROTOTYPE FEEDBACK LOOP
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("12. Prototype Feedback Loop", level=1)

doc.add_paragraph(
    "AIOS generates working prototypes from the BRD and manages iterative feedback. Before "
    "planning code changes, the retrieval system searches for original requirements evidence "
    'per feature. When a stakeholder says "the dashboard doesn\u2019t show what I expected," '
    "the system retrieves discovery chunks about dashboards \u2014 distinguishing \"we built it "
    'wrong" (original requirement was clear) from "the client evolved their thinking" '
    "(scope refinement). Different root causes produce different code strategies."
)

doc.add_paragraph(
    "Convergence tracking monitors how feedback changes between sessions. Increasing positive "
    "feedback means convergence. New objections mean divergence. The BRD updates with feedback "
    "evidence linked alongside original discovery evidence."
)

# ═══════════════════════════════════════════════════════════════════════════
# 13. COST ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("13. Cost Architecture", level=1)

doc.add_paragraph(
    "Every architectural decision optimizes for cost without sacrificing quality. The system "
    "uses a tiered model strategy: Haiku for structured extraction and classification, Sonnet "
    "for reasoning and synthesis, Opus for complex generation."
)

doc.add_heading("Pipeline Cost Breakdown", level=2)

add_table(
    ["Stage", "Model", "Calls", "Cost", "Time"],
    [
        ["Triage + classify", "Heuristic", "0", "Free", "<100ms"],
        ["Chunking", "Structural", "0", "Free", "<500ms"],
        ["Meta-tag enrichment", "Haiku \u00d7N parallel", "11", "~$0.011", "~3s"],
        ["Embedding (chunks)", "text-embedding-3-small", "11 batch", "~$0.001", "~1s"],
        ["Entity extraction", "Haiku \u00d7N parallel", "11", "~$0.003", "~5s"],
        ["Dedup (exact + fuzzy + Cohere + embedding)", "Python + Cohere + embedding", "~5", "~$0.001", "<100ms"],
        ["Speaker resolution", "Python (fuzzy match)", "0", "Free", "<10ms"],
        ["Entity embedding sync", "text-embedding-3-small", "1 batch", "~$0.001", "<500ms"],
        ["Memory Watcher", "Haiku \u00d7N parallel", "11", "~$0.002", "~3s"],
        ["Memory Synthesizer", "Sonnet (conditional)", "0\u20131", "~$0.02", "~4s"],
        ["Question auto-resolution", "Haiku (conditional)", "0\u20131", "~$0.001", "~1s"],
        ["Total per document", "", "", "~$0.04", "~15s wall clock"],
    ]
)

doc.add_heading("Retrieval Cost Breakdown", level=2)

add_table(
    ["Operation", "Cost", "Latency", "When"],
    [
        ["Retrieval: simple (1 round, no decompose)", "~$0.0005", "~300ms", "Simple chat questions"],
        ["Retrieval: standard (2 rounds, decompose + rerank)", "~$0.002", "~2.5s", "Complex chat, enrichment"],
        ["Retrieval: thorough (3 rounds, full agentic)", "~$0.003", "~4s", "Briefing, solution flow"],
        ["Cohere reranking", "~$0.0001", "~100ms", "Retrieval stage 3 (when >10 chunks)"],
        ["Graph neighborhood query", "Free", "~50ms", "Retrieval stage 2.5 (entity expansion)"],
        ["Relationship path finding", "Free", "~50ms", '"How does X relate to Y?" questions'],
        ["Tension detection", "Free", "~100ms", "Briefing, gap intelligence"],
    ]
)

add_callout(
    "Vector search, graph queries, and meta-tag filtering are PostgreSQL operations \u2014 free at query time. "
    "Meta-tags are computed once at ingestion and queried forever after. The expensive operations "
    "(LLM calls) happen only when new information enters. A project with 50 signals costs the same "
    "to query as a project with 5. The system gets more valuable without getting more expensive.",
    "Cost Scales with Value, Not Volume"
)

# ═══════════════════════════════════════════════════════════════════════════
# 14. ARCHITECTURAL ADVANTAGES
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("14. Architectural Advantages", level=1)

doc.add_heading("Everything in PostgreSQL", level=2)
doc.add_paragraph(
    "AIOS runs on a single Supabase instance: relational data, vector search (pgvector), "
    "JSONB metadata queries, graph traversal via signal_impact co-occurrence, real-time "
    "subscriptions, row-level security, edge functions, and file storage. No Pinecone. No "
    "Elasticsearch. No Redis. No Neo4j. One database, one deployment, one operational surface."
)

doc.add_heading("One Retrieval System, Every Flow", level=2)
doc.add_paragraph(
    "Every component that needs context calls the same retrieval function with different "
    "parameters. This means retrieval improvements \u2014 better reranking, richer meta-tags, "
    "graph expansion, new entity types \u2014 automatically benefit every flow in the platform. "
    "When we added Cohere reranking and graph expansion, chat, briefing, unlocks, gap intel, "
    "stakeholder profiling, and prototype planning all got better instantly."
)

doc.add_heading("Parallel-First Processing", level=2)
doc.add_paragraph(
    "The chunk is the unit of parallelism. Every LLM-dependent stage fans out one call per "
    "chunk and gathers results. This pattern \u2014 map individual chunks, reduce merged results "
    "\u2014 applies identically to extraction, meta-tagging, memory observation, and enrichment. "
    "Prompt caching makes this economic: all parallel calls share the same cached system prefix."
)

doc.add_heading("Adaptive Intelligence", level=2)
doc.add_paragraph(
    "The Pulse Engine means the system adapts to where each project actually is. Dedup "
    "thresholds shift based on entity saturation. Entity targets scale with signal velocity. "
    "Extraction directives change based on health scores. Stage transitions gate on hard "
    "criteria. The system isn\u2019t running the same static logic on every project \u2014 it\u2019s "
    "responding to the actual state of discovery."
)

doc.add_heading("Evidence as a First-Class Citizen", level=2)
doc.add_paragraph(
    "Every entity, every belief, every recommendation has a provenance chain. The extraction "
    "prompt requires evidence references. The dedup merge preserves evidence from all duplicates. "
    "The memory system links beliefs to facts to chunks. The retrieval system returns evidence "
    "alongside entities and beliefs. For consultants, this means defensible deliverables. For "
    "clients, this means trust."
)

doc.add_heading("Model-Agnostic by Design", level=2)
doc.add_paragraph(
    "The architecture separates \"what work needs to be done\" from \"which model does it.\" "
    "Extraction is structured classification \u2014 Haiku. Synthesis requires reasoning \u2014 Sonnet. "
    "Complex generation \u2014 Opus. Reranking \u2014 Cohere. Each stage specifies the capability "
    "required, not the specific model. When a faster or cheaper model ships, the swap is a "
    "single constant change."
)

# ═══════════════════════════════════════════════════════════════════════════
# 15. TECHNICAL STACK
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("15. Technical Stack", level=1)

add_table(
    ["Component", "Technology", "Why"],
    [
        ["Database & Auth", "Supabase (PostgreSQL 15)", "pgvector + JSONB + RLS + real-time"],
        ["Vector Search", "pgvector (HNSW index)", "No external vector DB; consistent with source data"],
        ["Hybrid Search", "JSONB + GIN index", "Meta-tag filtering without Elasticsearch"],
        ["Embeddings", "OpenAI text-embedding-3-small", "Best quality/cost at 1,536 dimensions"],
        ["Extraction & Tagging", "Claude Haiku 4.5", "Structured output, 10x faster and 4x cheaper than Sonnet"],
        ["Reasoning & Synthesis", "Claude Sonnet 4.6", "Tool use, complex reasoning, belief synthesis"],
        ["Complex Generation", "Claude Opus 4.6", "Prototype generation, strategic analysis"],
        ["Reranking", "Cohere rerank-v3.5", "Purpose-built for relevance scoring; Haiku fallback"],
        ["Semantic Dedup", "Cohere rerank-v3.5 + pgvector", "Tier 2.5 catches meaning equivalence"],
        ["Backend", "FastAPI (Python 3.11)", "Async-native, type-safe, minimal overhead"],
        ["Workflow Orchestration", "LangGraph", "Stateful, conditional routing, built-in persistence"],
        ["Similarity Matching", "RapidFuzz", "Sub-millisecond fuzzy matching for entity dedup"],
        ["Frontend", "React 18 + TypeScript + Next.js 14", "Consultant workbench and client portal"],
        ["Deployment", "Railway + Supabase Cloud", "Managed infrastructure, horizontal scaling ready"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════
# 16. IMPLEMENTATION STATUS
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading("16. Implementation Status", level=1)

add_table(
    ["#", "Capability", "Status"],
    [
        ["1", "Parallel chunk extraction (map-reduce Haiku)", "Shipped"],
        ["2", "Vector foundation (entity embeddings, memory embeddings, match RPCs)", "Shipped"],
        ["3", "Haiku meta-tag enrichment per chunk", "Shipped"],
        ["4", "Four-tier entity deduplication (exact + fuzzy + Cohere + embedding)", "Shipped"],
        ["5", "Speaker \u2192 stakeholder resolution", "Shipped"],
        ["6", "Chunk-fed memory watcher with semantic contradiction detection", "Shipped"],
        ["7", "Unified retrieval system (decompose \u2192 retrieve \u2192 graph expand \u2192 rerank \u2192 evaluate)", "Shipped"],
        ["8", "Chat vector pre-fetch (one-turn evidence answers)", "Shipped"],
        ["9", "Graph+vector hybrid retrieval (entity neighborhood expansion)", "Shipped"],
        ["10", "Cohere reranking with Haiku fallback", "Shipped"],
        ["11", "Solution flow, briefing, stakeholder intel \u2192 retrieval-grounded", "Shipped"],
        ["12", "Unlocks \u2192 parallel generation with evidence", "Shipped"],
        ["13", "Prototype updater \u2192 discovery context injection", "Shipped"],
        ["14", "Client portal confirmation clustering + dynamic question management", "Shipped"],
        ["15", "Pulse Engine (adaptive rules engine, dynamic thresholds)", "Shipped"],
        ["16", "Prompt caching across 20+ LLM call sites", "Shipped"],
        ["17", "Page-aware tool filtering + solution flow context layers", "Shipped"],
        ["18", "Question auto-resolution (signal + portal + cluster triggers)", "Shipped"],
        ["19", "Convergence tracking across prototype sessions", "Shipped"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════
# CLOSING
# ═══════════════════════════════════════════════════════════════════════════

doc.add_paragraph("")

add_callout(
    "AIOS doesn\u2019t generate documents \u2014 it builds a living intelligence layer between raw "
    "discovery and structured requirements. Every paragraph is indexed and tagged. Every entity "
    "has provenance and a vector. Every belief traces to evidence. A unified retrieval system "
    "connects it all \u2014 any component that needs context gets chunks, entities, beliefs, and "
    "graph relationships in one call. The Pulse Engine adapts system behavior to where each "
    "project actually is. The architecture is parallel-first, evidence-grounded, adaptive, and "
    "cost-optimized at every stage.",
    "The Bottom Line"
)

# ── Save ────────────────────────────────────────────────────────────────────
output_path = os.path.expanduser("~/Downloads/AIOS_Intelligence_Architecture_v3.docx")
doc.save(output_path)
print(f"Saved to: {output_path}")
